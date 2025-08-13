from flask import render_template, redirect, url_for, flash, send_file, request
from flask_login import login_user, login_required, logout_user, current_user
from datetime import datetime, date, timedelta
from werkzeug.security import check_password_hash, generate_password_hash
from io import BytesIO
from docx import Document
from openpyxl import Workbook

from app import app, db
from models import User, Patient, Appointment, MedicalRecord
from forms import LoginForm, NovoPacienteForm, MedicalRecordForm, AppointmentForm, UserForm
from decorators import role_required, roles_required


# --- Rotas de Autenticação e Páginas Principais ---

@app.route('/')
def index():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    return redirect(url_for('login'))

@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    form = LoginForm()
    if form.validate_on_submit():
        user = User.query.filter_by(username=form.username.data).first()
        # Corrigido nome do campo senha do form (ex: form.password ou form.senha)
        if user and check_password_hash(user.senha, form.password.data):  
            login_user(user)
            flash(f'Bem-vindo(a) de volta, {user.nome_completo}!', 'success')
            return redirect(url_for('dashboard'))
        else:
            flash('Usuário ou senha incorretos. Tente novamente.', 'danger')
    return render_template('login.html', form=form)

@app.route('/logout')
@login_required
def logout():
    logout_user()
    flash('Você saiu do sistema.', 'info')
    return redirect(url_for('login'))

# --- Usuários ---

@app.route('/novo_usuario', methods=['GET', 'POST'])
@login_required
@roles_required('administrador', 'gerencia')
def novo_usuario():
    form = UserForm()
    if form.validate_on_submit():
        if User.query.filter_by(username=form.username.data).first():
            flash('Usuário já existe com esse nome de usuário.', 'warning')
            return render_template('novo_usuario.html', form=form)
        novo_user = User(
            username=form.username.data,
            senha=generate_password_hash(form.senha.data),
            nome_completo=form.nome_completo.data,
            funcao=form.funcao.data
        )
        db.session.add(novo_user)
        db.session.commit()
        flash('Usuário criado com sucesso!', 'success')
        return redirect(url_for('dashboard'))
    return render_template('novo_usuario.html', form=form)


# --- Dashboard ---

@app.route('/dashboard')
@login_required
def dashboard():
    proximos_agendamentos = Appointment.query.filter(Appointment.data_hora >= datetime.utcnow())\
                                             .order_by(Appointment.data_hora.asc())\
                                             .limit(5).all()
    total_pacientes = Patient.query.count()
    pacientes = Patient.query.order_by(Patient.nome_completo).all()
    return render_template('dashboard.html', appointments=proximos_agendamentos,
                           total_pacientes=total_pacientes, pacientes=pacientes, show_flash=True)

@app.route('/dashboard_detalhado')
@login_required
def dashboard_detalhado():
    return render_template('dashboard_detalhado.html')

# --- Pacientes ---

@app.route('/novo_paciente', methods=['GET', 'POST'])
@login_required
def novo_paciente():
    form = NovoPacienteForm()
    if form.validate_on_submit():
        # Aqui criar instância do modelo Patient e salvar no banco
        paciente = Patient(
            nome_completo=form.nome_completo.data,
            nome_social=form.nome_social.data,
            idade=form.idade.data,
            data_nascimento=form.data_nascimento.data,
            paciente_dependente=form.paciente_dependente.data,
            nome_dependente=form.nome_dependente.data,
            data_nascimento_dependente=form.data_nascimento_dependente.data,
            idade_dependente=form.idade_dependente.data,
            endereco=form.endereco.data,
            email=form.email.data,
            telefone=form.telefone.data,
            escolaridade=form.escolaridade.data,
            religiao=form.religiao.data,
            estado_civil=form.estado_civil.data,
            servico_buscado=form.servico_buscado.data,
            servico_buscado_outro=form.servico_buscado_outro.data
        )
        db.session.add(paciente)
        db.session.commit()

        flash(f"Paciente {paciente.nome_completo} cadastrado com sucesso!", "success")
        return redirect(url_for('lista_pacientes'))

    # Calcula idade no GET se data de nascimento já preenchida
    if request.method == "GET" and form.data_nascimento.data:
        nascimento = form.data_nascimento.data
        hoje = date.today()
        idade_calc = hoje.year - nascimento.year - ((hoje.month, hoje.day) < (nascimento.month, nascimento.day))
        form.idade.data = idade_calc

    return render_template('novo_paciente.html', form=form)


@app.route('/lista_pacientes')
@login_required
def lista_pacientes():
    pacientes = Patient.query.order_by(Patient.nome_completo).all()
    return render_template('lista_pacientes.html', pacientes=pacientes)


# --- Prontuário ---

@app.route('/prontuario/<int:paciente_id>', methods=['GET', 'POST'])
@login_required
@roles_required('médico')
def prontuario(paciente_id):
    paciente = Patient.query.get_or_404(paciente_id)
    form = MedicalRecordForm()
    if form.validate_on_submit():
        registro = MedicalRecord(
            paciente_id=paciente_id,
            medico_id=current_user.id,
            data_sessao=form.data_sessao.data,
            evolucao=form.evolucao.data
        )
        db.session.add(registro)
        db.session.commit()
        flash('Evolução adicionada com sucesso!', 'success')
        return redirect(url_for('prontuario', paciente_id=paciente_id))

    evolucoes = MedicalRecord.query.filter_by(paciente_id=paciente_id).order_by(MedicalRecord.data_sessao.asc()).all()
    labels = [e.data_sessao.strftime('%d/%m/%Y') for e in evolucoes]
    valores = [len(e.evolucao) for e in evolucoes]

    return render_template('prontuario.html', paciente=paciente, evolucoes=evolucoes, form=form,
                           labels=labels, valores=valores)

# --- Agendamento ---

from sqlalchemy import and_

@app.route('/agendamento', methods=['GET', 'POST'])
@login_required
def agendamento():
    form = AppointmentForm()
    form.paciente_id.choices = [(p.id, p.nome_completo) for p in Patient.query.order_by(Patient.nome_completo).all()]

    salas = ['Sala 1', 'Sala 2', 'Sala 3', 'Sala 4']
    ultimos_agendamentos = Appointment.query.order_by(Appointment.data_hora.desc()).limit(5).all()

    if form.validate_on_submit():
        novo_inicio = form.data_hora.data
        duracao_min = int(form.duracao.data)
        novo_fim = novo_inicio + timedelta(minutes=duracao_min)
        sala_escolhida = form.sala.data

        # Validação do horário permitido (09:00 - 17:00)
        if not (9 <= novo_inicio.hour < 17 or (novo_inicio.hour == 17 and novo_inicio.minute == 0)):
            flash('O horário deve ser entre 09:00 e 17:00.', 'danger')
            return render_template('agendamento.html', form=form, ultimos_agendamentos=ultimos_agendamentos, salas=salas)

        # Buscar agendamentos da sala no mesmo dia para verificar conflito
        inicio_dia = novo_inicio.replace(hour=0, minute=0, second=0, microsecond=0)
        fim_dia = novo_inicio.replace(hour=23, minute=59, second=59, microsecond=999999)

        agendamentos_no_dia = Appointment.query.filter(
            Appointment.sala == sala_escolhida,
            Appointment.data_hora >= inicio_dia,
            Appointment.data_hora <= fim_dia
        ).all()

        # Verifica se há conflito de horário
        conflito = False
        for ag in agendamentos_no_dia:
            ag_inicio = ag.data_hora
            ag_fim = ag_inicio + timedelta(minutes=ag.duracao)
            # Checa se os períodos se sobrepõem
            if novo_inicio < ag_fim and novo_fim > ag_inicio:
                conflito = True
                break

        if conflito:
            flash('A sala selecionada está ocupada nesse horário. Por favor, escolha outro horário ou sala.', 'danger')
            return render_template('agendamento.html', form=form, ultimos_agendamentos=ultimos_agendamentos, salas=salas)

        # Se passou nas validações, cria o agendamento
        agendamento = Appointment(
            paciente_id=form.paciente_id.data,
            medico_id=current_user.id,
            sala=sala_escolhida,
            data_hora=novo_inicio,
            duracao=duracao_min,
            observacoes=form.observacoes.data
        )
        db.session.add(agendamento)
        db.session.commit()
        flash('Agendamento realizado com sucesso!', 'success')
        return redirect(url_for('agendamento'))

    return render_template('agendamento.html', form=form, ultimos_agendamentos=ultimos_agendamentos, salas=salas)

@app.route('/lista_agendamentos')
@login_required
def lista_agendamentos():
    agendamentos = Appointment.query.filter_by(medico_id=current_user.id).order_by(Appointment.data_hora).all()
    return render_template('lista_agendamentos.html', agendamentos=agendamentos, show_flash=False)

# --- Exportação de Documentos ---

@app.route('/exportar_docx/<int:paciente_id>')
@login_required
@role_required('médico')
def exportar_docx(paciente_id):
    paciente = Patient.query.get_or_404(paciente_id)
    evolucoes = MedicalRecord.query.filter_by(paciente_id=paciente_id).order_by(MedicalRecord.data_sessao).all()

    doc = Document()
    doc.add_heading(f'Prontuário de {paciente.nome_completo}', 0)
    for evo in evolucoes:
        doc.add_paragraph(f'Data: {evo.data_sessao.strftime("%d/%m/%Y")}')
        doc.add_paragraph(f'Evolução: {evo.evolucao}')
        doc.add_paragraph('---')

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    filename = f'prontuario_{paciente.nome_completo.replace(" ", "_")}.docx'
    return send_file(file_stream, as_attachment=True, download_name=filename,
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


@app.route('/exportar_xlsx/<int:paciente_id>')
@login_required
@role_required('médico')
def exportar_xlsx(paciente_id):
    paciente = Patient.query.get_or_404(paciente_id)
    evolucoes = MedicalRecord.query.filter_by(paciente_id=paciente_id).order_by(MedicalRecord.data_sessao).all()

    wb = Workbook()
    ws = wb.active
    ws.title = "Prontuário"
    ws.append(["Data", "Evolução"])
    for evo in evolucoes:
        ws.append([evo.data_sessao.strftime("%d/%m/%Y"), evo.evolucao])

    file_stream = BytesIO()
    wb.save(file_stream)
    file_stream.seek(0)
    filename = f'prontuario_{paciente.nome_completo.replace(" ", "_")}.xlsx'
    return send_file(file_stream, as_attachment=True, download_name=filename,
                     mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')

# --- Tratamento de erros ---

@app.errorhandler(404)
def not_found_error(error):
    return render_template('404.html'), 404

@app.errorhandler(500)
def internal_error(error):
    db.session.rollback()
    return render_template('500.html'), 500
